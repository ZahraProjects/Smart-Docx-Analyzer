import re
from docx import Document

# Function to reverse numbers
def reverse_number(num):
    return num[::-1]  # Reverse the number

# Function to reverse numbers in the text
def reverse_numbers_in_text(text):
    # Regular expression pattern to identify integers and decimals in both Persian and English formats
    pattern = r"[-+]?\d+[\.,٫]?\d*"
    
    # Replace Persian numbers with English in the text
    persian_to_english = str.maketrans("۰۱۲۳۴۵۶۷۸۹", "0123456789")

    # Using the pattern to find numbers and reverse them
    def replace_function(match):
        num = match.group(0)  # Extract the number from the text
        reversed_num = reverse_number(num.translate(persian_to_english).replace("٫", ".").replace(",", ""))
        return reversed_num  # Return the reversed number

    # Using the replace function to reverse numbers in the text
    new_text = re.sub(pattern, replace_function, text)
    
    return new_text

# Load text from a Word file
def load_text_from_docx(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + " "
    return text.strip()

# Save new text with reversed numbers to a new Word file
def save_new_text_to_docx(new_text, output_file_path):
    doc = Document()
    doc.add_paragraph(new_text)
    doc.save(output_file_path)

# Get text from Word file
input_file_path = "Smart-Docx-Analyzer\AI_Text_EN.docx"  # Path to the input Word file
text = load_text_from_docx(input_file_path)

# Reverse numbers
new_text = reverse_numbers_in_text(text)

# Save new text to Word file
output_file_path = "Smart-Docx-Analyzer\\new_text_with_reversed_numbers.docx"  # Path to the output Word file
save_new_text_to_docx(new_text, output_file_path)

print("New text with reversed numbers saved in new Word file.")